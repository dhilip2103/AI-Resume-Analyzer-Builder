from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.template.loader import render_to_string
import pdfkit
from docx import Document
import tempfile

# Step 1: Resume template selection
def resume_selection(request):
    return render(request, 'resume_builder/resume_selection.html')

# Step 2: Resume editor (based on selected template)
def resume_editor(request):
    if request.method == 'POST':
        selected_template = request.POST.get('template')
        valid_templates = ['resume_template1', 'resume_template2', 'resume_template3']
        
        if selected_template not in valid_templates:
            return redirect('resume_selection')
            
        return render(request, 'resume_builder/workspace.html', {
            # Correct path pattern
            'template_path': f'resume_builder/{selected_template}.html'
        })
    return redirect('resume_selection')

# Step 3: Preview the resume before download
def resume_preview(request):
    if request.method == 'POST':
        resume_data = request.POST  # Capture form data
        return render(request, 'resume_builder/resume_preview.html', {'data': resume_data})
    return redirect('resume_editor')

# Step 4: Export logic (PDF or DOCX)
def download_resume(request, format):
    if request.method == 'POST':
        resume_data = request.POST

        if format == 'pdf':
            html_string = render_to_string('resume_builder/resume_preview.html', {'data': resume_data})
            pdf_file = pdfkit.from_string(html_string, False)
            response = HttpResponse(pdf_file, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename=resume.pdf'
            return response

        elif format == 'docx':
            doc = Document()
            doc.add_heading(resume_data.get('name'), 0)
            doc.add_paragraph(f"Email: {resume_data.get('email')}")
            doc.add_paragraph(f"Phone: {resume_data.get('phone')}")
            # Add more content...
            with tempfile.NamedTemporaryFile(delete=True) as tmp:
                doc.save(tmp.name)
                tmp.seek(0)
                response = HttpResponse(tmp.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=resume.docx'
                return response

    return redirect('resume_preview')
